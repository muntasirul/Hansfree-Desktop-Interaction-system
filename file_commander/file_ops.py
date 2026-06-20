"""
file_ops.py — Executes file operations based on Gemini-parsed action JSON.
Supported actions: copy, move, delete, list, rename, create_folder, open
"""

import os
import shutil
import glob
import fnmatch
from pathlib import Path
from typing import Dict, Any, List, Tuple

# Image extensions for "all images" type requests
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif',
                    '.bmp', '.webp', '.tiff', '.svg', '.ico', '.heic'}
VIDEO_EXTENSIONS = {'.mp4', '.mkv', '.avi', '.mov', '.wmv', '.flv', '.webm'}
AUDIO_EXTENSIONS = {'.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma', '.m4a'}
DOC_EXTENSIONS = {'.pdf', '.doc', '.docx', '.xls',
                  '.xlsx', '.ppt', '.pptx', '.txt', '.csv'}


def resolve_filter(filter_str: str) -> set:
    """Convert filter string like '*.jpg,*.png' or 'images' to a set of extensions."""
    if not filter_str or filter_str.lower() in ('*', 'all', ''):
        return set()  # empty = all files

    mapping = {
        'images': IMAGE_EXTENSIONS,
        'image': IMAGE_EXTENSIONS,
        'photos': IMAGE_EXTENSIONS,
        'videos': VIDEO_EXTENSIONS,
        'video': VIDEO_EXTENSIONS,
        'audio': AUDIO_EXTENSIONS,
        'music': AUDIO_EXTENSIONS,
        'documents': DOC_EXTENSIONS,
        'docs': DOC_EXTENSIONS,
    }
    if filter_str.lower() in mapping:
        return mapping[filter_str.lower()]

    # Parse comma-separated globs like *.jpg,*.png
    exts = set()
    for part in filter_str.split(','):
        part = part.strip()
        if part.startswith('*.'):
            exts.add(part[1:].lower())  # keep the dot: .jpg
        elif part.startswith('.'):
            exts.add(part.lower())
        else:
            exts.add('.' + part.lower())
    return exts


def matches_filter(filename: str, ext_set: set) -> bool:
    if not ext_set:
        return True
    return Path(filename).suffix.lower() in ext_set


def get_matching_files(folder: str, ext_set: set) -> List[Path]:
    p = Path(folder)
    if not p.exists() or not p.is_dir():
        return []
    return [f for f in p.iterdir() if f.is_file() and matches_filter(f.name, ext_set)]


# ─────────────────────────────────────────────
#  Action handlers
# ─────────────────────────────────────────────

def do_copy(action: Dict) -> Tuple[bool, str]:
    src = action.get('source', '').strip()
    dst = action.get('destination', '').strip()
    filter_str = action.get('filter', '*')

    if not src or not dst:
        return False, "❌ Source or destination path is missing."

    src_path = Path(src)
    dst_path = Path(dst)

    if not src_path.exists():
        return False, f"❌ Source path does not exist: {src}"

    dst_path.mkdir(parents=True, exist_ok=True)
    ext_set = resolve_filter(filter_str)

    if src_path.is_file():
        shutil.copy2(src_path, dst_path / src_path.name)
        return True, f"✅ Copied '{src_path.name}' → '{dst}'"

    files = get_matching_files(src, ext_set)
    if not files:
        return False, f"⚠️ No matching files found in '{src}'."

    copied = []
    for f in files:
        shutil.copy2(f, dst_path / f.name)
        copied.append(f.name)

    return True, f"✅ Copied {len(copied)} file(s) to '{dst}':\n  " + "\n  ".join(copied)


def do_move(action: Dict) -> Tuple[bool, str]:
    src = action.get('source', '').strip()
    dst = action.get('destination', '').strip()
    filter_str = action.get('filter', '*')

    if not src or not dst:
        return False, "❌ Source or destination path is missing."

    src_path = Path(src)
    dst_path = Path(dst)

    if not src_path.exists():
        return False, f"❌ Source path does not exist: {src}"

    dst_path.mkdir(parents=True, exist_ok=True)
    ext_set = resolve_filter(filter_str)

    if src_path.is_file():
        shutil.move(str(src_path), dst_path / src_path.name)
        return True, f"✅ Moved '{src_path.name}' → '{dst}'"

    files = get_matching_files(src, ext_set)
    if not files:
        return False, f"⚠️ No matching files found in '{src}'."

    moved = []
    for f in files:
        shutil.move(str(f), dst_path / f.name)
        moved.append(f.name)

    return True, f"✅ Moved {len(moved)} file(s) to '{dst}':\n  " + "\n  ".join(moved)


def do_delete(action: Dict) -> Tuple[bool, str]:
    src = action.get('source', '').strip()
    filter_str = action.get('filter', '*')

    if not src:
        return False, "❌ Source path is missing."

    src_path = Path(src)
    if not src_path.exists():
        return False, f"❌ Path does not exist: {src}"

    ext_set = resolve_filter(filter_str)

    if src_path.is_file():
        src_path.unlink()
        return True, f"✅ Deleted file '{src}'"

    if not ext_set and filter_str in ('*', 'all', ''):
        # Delete the folder itself
        shutil.rmtree(src_path)
        return True, f"✅ Deleted folder and all contents: '{src}'"

    files = get_matching_files(src, ext_set)
    if not files:
        return False, f"⚠️ No matching files found in '{src}'."

    deleted = []
    for f in files:
        f.unlink()
        deleted.append(f.name)

    return True, f"✅ Deleted {len(deleted)} file(s) from '{src}':\n  " + "\n  ".join(deleted)


def do_list(action: Dict) -> Tuple[bool, str]:
    src = action.get('source', '').strip()
    filter_str = action.get('filter', '*')

    if not src:
        return False, "❌ Path is missing."

    src_path = Path(src)
    if not src_path.exists():
        return False, f"❌ Path does not exist: {src}"

    ext_set = resolve_filter(filter_str)

    if src_path.is_file():
        return True, f"📄 File: {src_path.name}  ({src_path.stat().st_size} bytes)"

    items = list(src_path.iterdir())
    if not items:
        return True, f"📁 Folder '{src}' is empty."

    lines = [f"📁 Contents of '{src}':"]
    dirs = [i for i in items if i.is_dir()]
    files = [i for i in items if i.is_file(
    ) and matches_filter(i.name, ext_set)]

    for d in sorted(dirs):
        lines.append(f"  📂 {d.name}/")
    for f in sorted(files):
        size = f.stat().st_size
        size_str = f"{size/1024:.1f} KB" if size >= 1024 else f"{size} B"
        lines.append(f"  📄 {f.name}  ({size_str})")

    lines.append(f"\n  {len(dirs)} folder(s), {len(files)} file(s)")
    return True, "\n".join(lines)


def do_rename(action: Dict) -> Tuple[bool, str]:
    src = action.get('source', '').strip()
    new_name = action.get('new_name', '').strip()

    if not src or not new_name:
        return False, "❌ Source path or new name is missing."

    src_path = Path(src)
    if not src_path.exists():
        return False, f"❌ Path does not exist: {src}"

    dst_path = src_path.parent / new_name
    src_path.rename(dst_path)
    return True, f"✅ Renamed '{src_path.name}' → '{new_name}'"


def do_create_folder(action: Dict) -> Tuple[bool, str]:
    dst = action.get('destination', '').strip(
    ) or action.get('source', '').strip()
    if not dst:
        return False, "❌ Folder path is missing."
    Path(dst).mkdir(parents=True, exist_ok=True)
    return True, f"✅ Created folder: '{dst}'"


def do_create_file(action: Dict) -> Tuple[bool, str]:
    dst = action.get('destination', '').strip(
    ) or action.get('source', '').strip()
    if not dst:
        return False, "❌ File path is missing."

    path = Path(dst)

    # Check if the AI only gave a directory or a trailing slash
    if path.is_dir() or dst.endswith(('/', '\\')):
        return False, f"❌ Path '{dst}' is a directory. Please provide a full file name (e.g., 'D:/Groq/document.docx')."

    # Ensure parents exist
    path.parent.mkdir(parents=True, exist_ok=True)

    # Create empty file
    path.touch()
    return True, f"✅ Created file: '{dst}'"


def do_write_file(action: Dict) -> Tuple[bool, str]:
    dst = action.get('destination', '').strip(
    ) or action.get('source', '').strip()
    content = action.get('content', '').strip()
    mode = action.get('mode', 'append').lower()  # 'append' or 'overwrite'

    if not dst:
        return False, "❌ File path is missing."
    if not content:
        return False, "❌ No content provided to write."

    path = Path(dst)
    ext = path.suffix.lower()

    try:
        if ext == '.docx':
            from docx import Document
            # If file exists and is not empty, try opening it
            try:
                if path.exists() and path.stat().st_size > 0 and mode == 'append':
                    doc = Document(dst)
                else:
                    doc = Document()
            except Exception:
                # If opening fails (corrupt or 0-byte), just start fresh
                doc = Document()

            doc.add_paragraph(content)
            doc.save(dst)
            return True, f"✍️ Wrote to Word document: '{dst}'"
        else:
            # Assume plain text for other extensions
            io_mode = 'a' if mode == 'append' else 'w'
            with open(dst, io_mode, encoding='utf-8') as f:
                f.write(content + "\n")
            return True, f"✍️ Wrote content to: '{dst}'"
    except PermissionError:
        return False, f"❌ Permission denied: The file '{dst}' is likely open in another program (e.g., Microsoft Word). Please close it and try again."
    except Exception as e:
        return False, f"❌ Failed to write to file: {e}"


def do_launch_app(action: Dict) -> Tuple[bool, str]:
    app_name = action.get('new_name', '').strip(
    ) or action.get('source', '').strip()
    if not app_name:
        return False, "❌ Application name is missing."

    # Mapping for common app names
    common_apps = {
        'chrome': 'chrome',
        'microsoft edge': 'msedge',  # Add this
        'edge': 'msedge',           # Add this
        'command prompt': 'cmd',    # Add this
        'cmd': 'cmd',               # Add this
        'snipping tool': 'snippingtool',  # Add this
        'google chrome': 'chrome',
        'browser': 'chrome',
        'notepad': 'notepad',
        'calc': 'calc',
        'calculator': 'calc',
        'explorer': 'explorer',
        'word': 'winword',
        'excel': 'excel',
        'powerpoint': 'powerpnt',
        'paint': 'mspaint'
    }

    cmd_name = common_apps.get(app_name.lower(), app_name)

    try:
        import subprocess
        # Using 'start' to launch on Windows
        subprocess.Popen(f"start {cmd_name}", shell=True)
        return True, f"🚀 Launching application: '{app_name}'"
    except Exception as e:
        return False, f"❌ Failed to launch '{app_name}': {e}"


def do_close_app(action: Dict) -> Tuple[bool, str]:
    app_name = action.get('new_name', '').strip(
    ) or action.get('source', '').strip()
    if not app_name:
        return False, "❌ Application name is missing."

    name_low = app_name.lower()

    # 1. Safely close File Explorer windows/drives without killing the Windows desktop
    if "drive" in name_low or "explorer" in name_low or name_low.endswith(":/") or name_low.endswith(":\\"):
        try:
            import subprocess
            # PowerShell script to gracefully close open File Explorer folder windows
            ps_script = "(New-Object -comObject Shell.Application).Windows() | ForEach-Object { $_.Quit() }"
            subprocess.run(["powershell", "-Command",
                           ps_script], capture_output=True)
            return True, f"⏹️ Closed folder/drive windows."
        except Exception as e:
            return False, f"❌ Failed to close folders: {e}"

    # 2. Mapping for common process names
    common_procs = {
        'chrome': 'chrome.exe',
        'google chrome': 'chrome.exe',
        'browser': 'chrome.exe',
        'notepad': 'notepad.exe',
        'calc': 'CalculatorApp.exe',        # Updated to modern Windows Calculator
        'calculator': 'CalculatorApp.exe',
        'word': 'winword.exe',
        'microsoft word': 'winword.exe',
        'excel': 'excel.exe',
        'powerpoint': 'powerpnt.exe',
        'paint': 'mspaint.exe',
        'snipping tool': 'SnippingTool.exe',
        'vscode': 'Code.exe',
        'visual studio code': 'Code.exe',
        'vs code': 'Code.exe',
        'code': 'Code.exe',
        'spotify': 'Spotify.exe',
        'discord': 'Discord.exe',
        'slack': 'Slack.exe',
        'zoom': 'Zoom.exe',
        'teams': 'Teams.exe',
        'vlc': 'vlc.exe',
        'skype': 'Skype.exe'
    }

    proc_name = common_procs.get(name_low)

    try:
        import subprocess
        if proc_name:
            # Added double quotes around {proc_name} to handle spaces
            result = subprocess.run(
                f'taskkill /F /IM "{proc_name}"', shell=True, capture_output=True, text=True)
            if result.returncode == 0 or "successfully terminated" in result.stdout.lower():
                return True, f"⏹️ Closed application: '{app_name}'"
            else:
                return False, f"❌ Application '{app_name}' is not running."
        else:
            # Fallback execution, also wrapped in double quotes
            fallback_name = app_name if app_name.lower().endswith(
                '.exe') else f"{app_name}.exe"
            result = subprocess.run(
                f'taskkill /F /IM "{fallback_name}"', shell=True, capture_output=True, text=True)
            if result.returncode == 0 or "successfully terminated" in result.stdout.lower():
                return True, f"⏹️ Closed application: '{app_name}'"
            else:
                return False, f"❌ Could not close '{app_name}' - application may not be running."
    except Exception as e:
        return False, f"❌ Failed to close '{app_name}': {e}"


def do_open(action: Dict) -> Tuple[bool, str]:
    src = action.get('source', '').strip()
    if not src:
        return False, "❌ Path is missing."
    src_path = Path(src)
    if not src_path.exists():
        # Maybe it's an app?
        return do_launch_app(action)

    os.startfile(str(src_path))
    return True, f"✅ Opened: '{src}'"


# ─────────────────────────────────────────────
#  Dispatcher
# ─────────────────────────────────────────────

ACTION_MAP = {
    'copy': do_copy,
    'move': do_move,
    'delete': do_delete,
    'remove': do_delete,
    'list': do_list,
    'show': do_list,
    'rename': do_rename,
    'create_folder': do_create_folder,
    'mkdir': do_create_folder,
    'create_file': do_create_file,
    'touch': do_create_file,
    'write_file': do_write_file,
    'open': do_open,
    'launch_app': do_launch_app,
    'open_app': do_launch_app,
    'close': do_close_app,      # <-- Add this line
    'close_app': do_close_app,
    'exit_app': do_close_app,
    'stop_app': do_close_app,
}


def execute_action(action: Dict[str, Any]) -> Tuple[bool, str]:
    """Main entry point. Receives parsed action dict and runs the correct handler."""
    act = action.get('action', '').lower()
    handler = ACTION_MAP.get(act)
    if not handler:
        return False, f"❌ Unknown action: '{act}'. Supported: {', '.join(ACTION_MAP.keys())}"
    try:
        return handler(action)
    except PermissionError as e:
        return False, f"❌ Permission denied: {e}"
    except Exception as e:
        return False, f"❌ Error during '{act}': {e}"
